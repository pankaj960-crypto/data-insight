"""Export reports to PDF, Excel, Word, and CSV."""

import io
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from app.config import get_settings

settings = get_settings()


def _ensure_report_dir(user_id: int) -> Path:
    path = Path(settings.upload_dir) / str(user_id) / "reports"
    path.mkdir(parents=True, exist_ok=True)
    return path


def export_csv(df: pd.DataFrame, user_id: int, dataset_name: str) -> tuple[str, int]:
    report_dir = _ensure_report_dir(user_id)
    filename = f"{uuid.uuid4().hex}_{dataset_name}_processed.csv"
    file_path = report_dir / filename
    df.to_csv(file_path, index=False)
    return str(file_path), file_path.stat().st_size


def export_excel(
    df: pd.DataFrame,
    analysis: Dict[str, Any],
    user_id: int,
    dataset_name: str,
) -> tuple[str, int]:
    report_dir = _ensure_report_dir(user_id)
    filename = f"{uuid.uuid4().hex}_{dataset_name}_analysis.xlsx"
    file_path = report_dir / filename

    with pd.ExcelWriter(file_path, engine="openpyxl") as writer:
        df.to_excel(writer, sheet_name="Cleaned Data", index=False)
        summary_rows = [
            ["Metric", "Value"],
            ["Rows", analysis.get("rows", len(df))],
            ["Columns", analysis.get("columns", len(df.columns))],
            ["Quality Score", analysis.get("data_quality_score", "N/A")],
            ["Duplicates", analysis.get("duplicate_count", 0)],
        ]
        pd.DataFrame(summary_rows[1:], columns=summary_rows[0]).to_excel(
            writer, sheet_name="Summary", index=False
        )
        if analysis.get("insights"):
            pd.DataFrame({"Insight": analysis["insights"]}).to_excel(
                writer, sheet_name="Insights", index=False
            )
        missing = analysis.get("missing_values", {})
        if missing:
            pd.DataFrame(
                [{"Column": k, "Missing": v} for k, v in missing.items()]
            ).to_excel(writer, sheet_name="Missing Values", index=False)

    return str(file_path), file_path.stat().st_size


def export_pdf(
    df: pd.DataFrame,
    analysis: Dict[str, Any],
    insights: List[str],
    user_id: int,
    dataset_name: str,
    predictions: Optional[Dict[str, Any]] = None,
) -> tuple[str, int]:
    report_dir = _ensure_report_dir(user_id)
    filename = f"{uuid.uuid4().hex}_{dataset_name}_report.pdf"
    file_path = report_dir / filename

    doc = SimpleDocTemplate(str(file_path), pagesize=letter)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("Title", parent=styles["Heading1"], fontSize=18, spaceAfter=12)
    elements = []

    elements.append(Paragraph(f"DataInsight AI — {dataset_name}", title_style))
    elements.append(Paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}", styles["Normal"]))
    elements.append(Spacer(1, 0.3 * inch))

    summary_data = [
        ["Rows", str(analysis.get("rows", len(df)))],
        ["Columns", str(analysis.get("columns", len(df.columns)))],
        ["Quality Score", str(analysis.get("data_quality_score", "N/A"))],
        ["Duplicates", str(analysis.get("duplicate_count", 0))],
    ]
    table = Table([["Metric", "Value"]] + summary_data)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4F46E5")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ]
        )
    )
    elements.append(Paragraph("Dataset Summary", styles["Heading2"]))
    elements.append(table)
    elements.append(Spacer(1, 0.2 * inch))

    if insights:
        elements.append(Paragraph("Key Insights", styles["Heading2"]))
        for ins in insights[:10]:
            elements.append(Paragraph(f"• {ins}", styles["Normal"]))
        elements.append(Spacer(1, 0.2 * inch))

    if predictions:
        elements.append(Paragraph("Prediction Results", styles["Heading2"]))
        for k, v in predictions.get("metrics", {}).items():
            elements.append(Paragraph(f"• {k}: {v}", styles["Normal"]))

    doc.build(elements)
    return str(file_path), file_path.stat().st_size


def export_word(
    df: pd.DataFrame,
    analysis: Dict[str, Any],
    insights: List[str],
    user_id: int,
    dataset_name: str,
) -> tuple[str, int]:
    report_dir = _ensure_report_dir(user_id)
    filename = f"{uuid.uuid4().hex}_{dataset_name}_executive.docx"
    file_path = report_dir / filename

    doc = Document()
    doc.add_heading(f"Executive Summary — {dataset_name}", 0)
    doc.add_paragraph(f"Report generated on {datetime.utcnow().strftime('%B %d, %Y')}")

    doc.add_heading("Overview", level=1)
    doc.add_paragraph(
        f"This dataset contains {analysis.get('rows', len(df)):,} rows and "
        f"{analysis.get('columns', len(df.columns))} columns with a data quality score of "
        f"{analysis.get('data_quality_score', 'N/A')}/100."
    )

    doc.add_heading("Key Insights", level=1)
    for ins in (insights or [])[:10]:
        doc.add_paragraph(ins, style="List Bullet")

    doc.add_heading("Sample Data", level=1)
    table = doc.add_table(rows=1, cols=min(5, len(df.columns)))
    hdr = table.rows[0].cells
    for i, col in enumerate(df.columns[:5]):
        hdr[i].text = str(col)
    for _, row in df.head(5).iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns[:5]):
            cells[i].text = str(row[col])[:50]

    doc.save(file_path)
    return str(file_path), file_path.stat().st_size
